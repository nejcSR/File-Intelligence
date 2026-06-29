import os
import json
import pdfplumber
import pytesseract
import csv
import io
import docx
import platform
from PIL import Image
from fastapi import FastAPI, UploadFile, File, Depends, HTTPException, BackgroundTasks
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from dotenv import load_dotenv

from app.database import engine, get_db, Base, SessionLocal
from app.models import Document, Extraction
from app.extractor import extract_document

if platform.system() == "Windows":
    pytesseract.pytesseract.tesseract_cmd = r"F:\Tesseract\tesseract.exe"

load_dotenv()
Base.metadata.create_all(bind=engine)
app = FastAPI()

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

app.mount("/static", StaticFiles(directory="app/static"), name="static")


@app.get("/")
def serve_frontend():
    return FileResponse("app/static/index.html")


def process_document_heavy_lifting(doc_id: int, file_path: str):
    """Performs OCR and AI extraction in a background thread."""
    db: Session = SessionLocal()
    try:
        text = ""
        filename_lower = file_path.lower()

        if filename_lower.endswith(".pdf"):
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

        elif filename_lower.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif", ".jfif")):
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)

        elif filename_lower.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

        elif filename_lower.endswith(".docx"):
            document = docx.Document(file_path)
            text = "\n".join([para.text for para in document.paragraphs])

        doc = db.query(Document).filter(Document.id == doc_id).first()
        if not doc:
            return

        doc.raw_text = text

        result = extract_document(text)

        extraction = Extraction(
            document_id=doc.id,
            document_type=result.get("document_type"),
            extracted_fields=json.dumps(result.get("key_fields", {})),
            anomalies=json.dumps(result.get("anomalies", [])),
            confidence=result.get("confidence"),
            summary=result.get("summary")
        )
        db.add(extraction)

        doc.document_type = result.get("document_type")
        doc.status = "done"
        db.commit()

    except Exception as e:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if doc:
            doc.status = "failed"
            doc.failure_reason = str(e)
            db.commit()
        print(f"Background processing failed for doc {doc_id}: {str(e)}")


@app.post("/upload")
async def upload_file(
    background_tasks: BackgroundTasks,
    files: list[UploadFile] = File(...),
    db: Session = Depends(get_db)
):
    allowed_extensions = (
        ".pdf", ".txt",".docx",
        ".png", ".jpg", ".jpeg", ".webp", ".gif", ".jfif"
    )

    results = []

    for file in files:
        if not file.filename.lower().endswith(allowed_extensions):
            results.append({
                "filename": file.filename,
                "status": "rejected",
                "message": "Unsupported file type."
            })
            continue

        file_path = os.path.join(UPLOAD_DIR, file.filename)
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)

        doc = Document(
            filename=file.filename,
            status="processing",
            raw_text=""
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)

        background_tasks.add_task(process_document_heavy_lifting, doc.id, file_path)

        results.append({
            "document_id": doc.id,
            "filename": file.filename,
            "status": "processing",
            "message": "Queued for processing."
        })

    return results


@app.get("/documents")
def get_documents(db: Session = Depends(get_db)):
    docs = db.query(Document).order_by(Document.upload_date.desc()).all()
    result = []
    for doc in docs:
        extraction = db.query(Extraction).filter(Extraction.document_id == doc.id).first()
        
        anomalies_list = []
        if extraction and extraction.anomalies:
            try:
                anomalies_list = json.loads(extraction.anomalies)
            except Exception:
                anomalies_list = []

        result.append({
            "id": doc.id,
            "filename": doc.filename,
            "document_type": doc.document_type,
            "status": doc.status,
            "upload_date": doc.upload_date,
            "anomalies": anomalies_list,
            "confidence": extraction.confidence if extraction else None,
            "summary": extraction.summary if extraction else None
        })
    return result

@app.get("/documents/{doc_id}")
def get_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    extraction = db.query(Extraction).filter(Extraction.document_id == doc_id).first()

    return {
        "id": doc.id,
        "filename": doc.filename,
        "document_type": doc.document_type,
        "status": doc.status,
        "upload_date": doc.upload_date,
        "raw_text": doc.raw_text,
        "key_fields": json.loads(extraction.extracted_fields) if extraction and extraction.extracted_fields else {},
        "anomalies": json.loads(extraction.anomalies) if extraction and extraction.anomalies else [],
        "confidence": extraction.confidence if extraction else None,
        "summary": extraction.summary if extraction else None,
        "failure_reason": doc.failure_reason
    }


@app.get("/export/csv")
def export_csv(db: Session = Depends(get_db)):
    docs = db.query(Document).order_by(Document.upload_date.desc()).all()

    output = io.StringIO()
    writer = csv.writer(output)

    writer.writerow(["ID", "Filename", "Type", "Status", "Confidence", "Anomalies", "Summary", "Date"])

    for doc in docs:
        extraction = db.query(Extraction).filter(Extraction.document_id == doc.id).first()
        
        anomalies = []
        if extraction and extraction.anomalies:
            try:
                anomalies = json.loads(extraction.anomalies)
            except Exception:
                anomalies = []

        writer.writerow([
            doc.id,
            doc.filename,
            doc.document_type or "",
            doc.status,
            extraction.confidence if extraction else "",
            " | ".join(anomalies),
            extraction.summary if extraction else "",
            doc.upload_date
        ])

    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=documents.csv"}
    )


@app.delete("/documents/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    file_path = os.path.join(UPLOAD_DIR, doc.filename)
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception as e:
            print(f"Failed to delete disk storage file: {e}")

    extraction = db.query(Extraction).filter(Extraction.document_id == doc_id).first()
    if extraction:
        db.delete(extraction)

    db.delete(doc)
    db.commit()

    return {"message": f"Document {doc_id} deleted successfully"}

@app.delete("/documents")
def delete_all_documents(db: Session = Depends(get_db)):
    # 1. Clear out all extraction records
    db.query(Extraction).delete()
    
    # 2. Clear out all document rows
    db.query(Document).delete()
    db.commit()

    # 3. Clean up the physical uploads folder
    if os.path.exists(UPLOAD_DIR):
        for filename in os.listdir(UPLOAD_DIR):
            file_path = os.path.join(UPLOAD_DIR, filename)
            try:
                if os.path.isfile(file_path):
                    os.remove(file_path)
            except Exception as e:
                print(f"Failed to delete file {filename} from disk: {e}")

    return {"message": "All documents and metadata successfully wiped clean."}